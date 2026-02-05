"""
Вспомогательные функции для сборки полного ТЗ и экспорта в формат DOCX.

Основная идея:
- на уровне БД мы храним отдельные части ТЗ (разные разделы) в таблице GenerationLog;
- пользователь выбирает нужные записи (разделы) в веб-интерфейсе;
- здесь мы:
    * сортируем разделы в логичном порядке;
    * добавляем заголовки;
    * собираем единый документ;
    * формируем DOCX-файл (python-docx).

Использование (из Flask-вьюхи):
    from .docx_utils import build_docx_from_logs
    bytes_io = build_docx_from_logs(
        sections=selected_logs,
        doc_title="Техническое задание на разработку ИС учета сотрудников",
        project_name="Система учета сотрудников МУИВ",
        project_domain="кадровый учет в вузе",
        comment="Черновая версия для согласования"
    )
    return send_file(
        bytes_io,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name="TZ_hr_muiv.docx",
    )
"""

from __future__ import annotations

from io import BytesIO
from typing import Iterable, List, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from .models import GenerationLog


# --------------------------------------------------------------------------------------
# Карта "нормализованное название раздела" → (порядковый номер, отображаемый заголовок)
# --------------------------------------------------------------------------------------

SECTION_ORDER_MAP = {
    # Классическая тройка, которую мы сейчас генерируем
    "основания для разработки": (1, "1. Основания для разработки"),
    "назначение системы": (2, "2. Назначение системы"),
    "требования к системе": (3, "3. Требования к системе"),
    # При желании позже можно расширить:
    # "характеристика объектов автоматизации": (4, "4. Характеристика объектов автоматизации"),
    # и т.д.
}


def _normalize_section_name(name: str) -> str:
    """Приведение названия раздела к нижнему регистру и обрезке пробелов."""
    return (name or "").strip().lower()


def _get_order_and_title(raw_name: str) -> Tuple[int, str]:
    """
    Возвращает:
        - порядковый номер раздела в документе;
        - отображаемый заголовок (с номером и корректным регистром).

    Если раздел не известен SECTION_ORDER_MAP, он попадает "в хвост" (порядок 99),
    а в заголовок подставляется исходное название.
    """
    norm = _normalize_section_name(raw_name)
    if norm in SECTION_ORDER_MAP:
        order, title = SECTION_ORDER_MAP[norm]
        return order, title
    # по умолчанию — в конец, под оригинальным названием
    return 99, raw_name or "Раздел технического задания"


def _sort_logs_for_document(sections: Iterable[GenerationLog]) -> List[Tuple[int, str, str]]:
    """
    Преобразует список GenerationLog в отсортированный список кортежей:
        (order, title, text)

    order       — число для сортировки;
    title       — заголовок раздела (с номером, если есть в карте);
    text        — текст раздела.
    """
    prepared: List[Tuple[int, str, str]] = []

    for log in sections:
        section_name = log.section_name or ""
        generated_text = log.generated_text or ""
        order, title = _get_order_and_title(section_name)
        prepared.append((order, title, generated_text))

    # сортируем по order, а при равенстве — по названию
    prepared.sort(key=lambda x: (x[0], x[1]))

    return prepared


def _add_title_page(
    doc: Document,
    doc_title: str,
    project_name: str | None = None,
    project_domain: str | None = None,
    comment: str | None = None,
) -> None:
    """
    Добавляет простую титульную страницу в документ:
    - название документа;
    - информация о проекте;
    - при необходимости — краткое примечание.
    """
    # Основной заголовок
    title_par = doc.add_paragraph()
    title_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_par.add_run(doc_title or "Техническое задание")
    run.bold = True
    run.font.size = Pt(16)

    # Дополнительная информация о проекте
    if project_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Проект: {project_name}")
        run.font.size = Pt(12)

    if project_domain:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Предметная область: {project_domain}")
        run.font.size = Pt(12)

    if comment:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(comment)
        run.italic = True
        run.font.size = Pt(11)

    # Небольшой отступ и разрыв страницы
    doc.add_paragraph()
    doc.add_page_break()


def _add_section(doc: Document, title: str, text: str) -> None:
    """
    Добавляет в документ один раздел:
    - заголовок уровня Heading 1;
    - основной текст раздела, разбитый на параграфы по пустым строкам.
    """
    # Заголовок раздела
    doc.add_heading(title, level=1)

    # Основной текст: разбиваем по "двойным" переводам строки
    if not text:
        doc.add_paragraph("(раздел не заполнен)")
        return

    # Сначала нормализуем переносы
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    blocks = [b.strip() for b in normalized.split("\n\n") if b.strip()]

    if not blocks:
        doc.add_paragraph(text.strip())
        return

    for block in blocks:
        doc.add_paragraph(block)


def build_docx_from_logs(
    sections: Iterable[GenerationLog],
    doc_title: str,
    project_name: str | None = None,
    project_domain: str | None = None,
    comment: str | None = None,
) -> BytesIO:
    """
    Формирует DOCX-документ на основе выбранных записей GenerationLog.

    Параметры:
        sections      — итерируемый набор записей журнала (GenerationLog),
                        каждая запись содержит section_name и generated_text;
        doc_title     — заголовок документа (будет на титульной странице);
        project_name  — имя проекта (опционально, для титульного листа);
        project_domain— предметная область (опционально);
        comment       — дополнительное примечание (опционально).

    Возвращает:
        BytesIO с готовым DOCX-файлом (позиция указателя = 0).
    """
    # 1. Сортируем разделы в логичном порядке
    sorted_sections = _sort_logs_for_document(sections)

    # 2. Создаём документ
    doc = Document()

    # 3. Титульная страница
    _add_title_page(
        doc=doc,
        doc_title=doc_title,
        project_name=project_name,
        project_domain=project_domain,
        comment=comment,
    )

    # 4. Основные разделы
    for _, section_title, section_text in sorted_sections:
        _add_section(doc, section_title, section_text)

    # 5. Сохранение в память
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
